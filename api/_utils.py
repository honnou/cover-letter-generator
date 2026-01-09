"""
Shared utilities for serverless functions
"""

import os
import io
import PyPDF2
from docx import Document
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import anthropic
import openai

# Configuration
AI_PROVIDER = os.getenv('AI_PROVIDER', 'anthropic')
ANTHROPIC_API_KEY = os.getenv('ANTHROPIC_API_KEY', '')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', '')


def extract_text_from_pdf(file_stream):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")


def extract_text_from_docx(file_stream):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_stream)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")


def extract_text_from_file(file_data, filename):
    """Extract text from uploaded file based on file type"""
    filename_lower = filename.lower()

    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(io.BytesIO(file_data))
    elif filename_lower.endswith(('.doc', '.docx')):
        return extract_text_from_docx(io.BytesIO(file_data))
    elif filename_lower.endswith('.txt'):
        return file_data.decode('utf-8')
    else:
        raise Exception(f"Unsupported file type: {filename}")


def generate_cover_letter_with_ai(example_cover_letter, job_description, ai_provider='anthropic'):
    """Use AI to generate adapted cover letter"""

    prompt = f"""You are a professional career advisor. You need to adapt an existing cover letter for a new job opportunity.

EXAMPLE COVER LETTER:
{example_cover_letter}

NEW JOB DESCRIPTION:
{job_description}

Please create a new cover letter for the new job that:
1. Maintains the tone and style of the example cover letter
2. Highlights relevant skills and experiences that match the new job requirements
3. Customizes the content to address the specific role and company
4. Keeps the same general structure and format
5. Is professional, compelling, and tailored to the new position

Please output ONLY the new cover letter text, without any preamble or explanation."""

    try:
        if ai_provider == 'anthropic' and ANTHROPIC_API_KEY:
            # Initialize client here, not at module level
            client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            message = client.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=2000,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            return message.content[0].text

        elif ai_provider == 'openai' and OPENAI_API_KEY:
            # Initialize OpenAI client here
            openai.api_key = OPENAI_API_KEY
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a professional career advisor helping to write cover letters."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000
            )
            return response.choices[0].message.content

        else:
            # Fallback: Simple template-based generation (no AI)
            return f"""Dear Hiring Manager,

I am writing to express my strong interest in the position described in your job posting. Based on my background and experience, I believe I would be an excellent fit for this role.

{example_cover_letter.split('.')[-3] if '.' in example_cover_letter else 'My relevant experience'}

The requirements outlined in your job description align well with my skills and career goals. I am particularly excited about the opportunity to contribute to your team.

Thank you for considering my application. I look forward to discussing how I can contribute to your organization.

Sincerely,
[Your Name]

---
NOTE: This is a basic template. Configure ANTHROPIC_API_KEY or OPENAI_API_KEY for AI-generated cover letters.
"""
    except Exception as e:
        raise Exception(f"Error generating cover letter: {str(e)}")


def generate_docx(text):
    """Generate DOCX document from text"""
    doc = DocxDocument()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add paragraphs
    for line in text.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            # Format paragraphs
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
        else:
            doc.add_paragraph()  # Blank line

    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes


def generate_pdf(text):
    """Generate PDF document from text"""
    pdf_bytes = io.BytesIO()

    # Create PDF
    doc = SimpleDocTemplate(
        pdf_bytes,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    # Styles
    styles = getSampleStyleSheet()
    style = ParagraphStyle(
        'CustomStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=11,
        leading=14,
        spaceAfter=12
    )

    # Build content
    story = []
    for line in text.split('\n'):
        if line.strip():
            para = Paragraph(line.strip(), style)
            story.append(para)
        else:
            story.append(Spacer(1, 0.2 * inch))

    # Generate PDF
    doc.build(story)
    pdf_bytes.seek(0)
    return pdf_bytes
