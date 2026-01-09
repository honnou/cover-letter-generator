#!/usr/bin/env python3
"""
Cover Letter Generator Backend Server
Handles file uploads, AI processing, and document generation
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import io
import tempfile
from datetime import datetime
from pathlib import Path

# Document parsing
import PyPDF2
from docx import Document
import markdown2

# Document generation
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

# AI Integration
import anthropic
import openai

app = Flask(__name__)
CORS(app)

# Configuration
AI_PROVIDER = os.getenv('AI_PROVIDER', 'anthropic')  # 'anthropic' or 'openai'
ANTHROPIC_API_KEY = os.getenv('ANTHROPIC_API_KEY', '')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', '')

# Initialize AI clients
anthropic_client = None
openai_client = None

if ANTHROPIC_API_KEY:
    anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
if OPENAI_API_KEY:
    openai.api_key = OPENAI_API_KEY
    openai_client = openai


def extract_text_from_pdf(file_stream):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")


def extract_text_from_docx(file_stream):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_stream)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")


def extract_text_from_file(file):
    """Extract text from uploaded file based on file type"""
    filename = file.filename.lower()

    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file.stream)
    elif filename.endswith(('.doc', '.docx')):
        return extract_text_from_docx(file.stream)
    elif filename.endswith('.txt'):
        return file.read().decode('utf-8')
    else:
        raise Exception(f"Unsupported file type: {filename}")


def generate_cover_letter_with_ai(example_cover_letter, job_description, ai_provider='anthropic'):
    """Use AI to generate adapted cover letter"""

    prompt = f"""You are a professional career advisor. You need to adapt an existing cover letter for a new job opportunity.

EXAMPLE COVER LETTER:
{example_cover_letter}

NEW JOB DESCRIPTION:
{job_description}

Please create a new cover letter for the new job that:
1. Maintains the tone and style of the example cover letter
2. Highlights relevant skills and experiences that match the new job requirements
3. Customizes the content to address the specific role and company
4. Keeps the same general structure and format
5. Is professional, compelling, and tailored to the new position

Please output ONLY the new cover letter text, without any preamble or explanation."""

    try:
        if ai_provider == 'anthropic' and anthropic_client:
            message = anthropic_client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=2000,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            return message.content[0].text

        elif ai_provider == 'openai' and openai_client:
            response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a professional career advisor helping to write cover letters."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000
            )
            return response.choices[0].message.content

        else:
            # Fallback: Simple template-based generation (no AI)
            return f"""Dear Hiring Manager,

I am writing to express my strong interest in the position described in your job posting. Based on my background and experience, I believe I would be an excellent fit for this role.

{example_cover_letter.split('.')[-3] if '.' in example_cover_letter else 'My relevant experience'}

The requirements outlined in your job description align well with my skills and career goals. I am particularly excited about the opportunity to contribute to your team.

Thank you for considering my application. I look forward to discussing how I can contribute to your organization.

Sincerely,
[Your Name]

---
NOTE: This is a basic template. Please configure ANTHROPIC_API_KEY or OPENAI_API_KEY in your environment for AI-generated cover letters.
"""
    except Exception as e:
        raise Exception(f"Error generating cover letter: {str(e)}")


def generate_docx(text):
    """Generate DOCX document from text"""
    doc = DocxDocument()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add paragraphs
    for line in text.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            # Format paragraphs
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
        else:
            doc.add_paragraph()  # Blank line

    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes


def generate_pdf(text):
    """Generate PDF document from text"""
    pdf_bytes = io.BytesIO()

    # Create PDF
    doc = SimpleDocTemplate(
        pdf_bytes,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    # Styles
    styles = getSampleStyleSheet()
    style = ParagraphStyle(
        'CustomStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=11,
        leading=14,
        spaceAfter=12
    )

    # Build content
    story = []
    for line in text.split('\n'):
        if line.strip():
            para = Paragraph(line.strip(), style)
            story.append(para)
        else:
            story.append(Spacer(1, 0.2 * inch))

    # Generate PDF
    doc.build(story)
    pdf_bytes.seek(0)
    return pdf_bytes


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'ai_provider': AI_PROVIDER,
        'ai_configured': bool(ANTHROPIC_API_KEY or OPENAI_API_KEY)
    })


@app.route('/generate', methods=['POST'])
def generate_cover_letter():
    """Main endpoint to generate cover letter"""
    try:
        # Get form data
        example_text = request.form.get('example_text', '')
        job_description_text = request.form.get('job_description_text', '')

        # Get uploaded files
        example_file = request.files.get('example_file')
        job_description_file = request.files.get('job_description_file')

        # Extract text from files or use provided text
        if example_file and example_file.filename:
            example_text = extract_text_from_file(example_file)

        if job_description_file and job_description_file.filename:
            job_description_text = extract_text_from_file(job_description_file)

        # Validate inputs
        if not example_text or not job_description_text:
            return jsonify({
                'error': 'Both example cover letter and job description are required'
            }), 400

        # Generate new cover letter using AI
        new_cover_letter = generate_cover_letter_with_ai(
            example_text,
            job_description_text,
            AI_PROVIDER
        )

        return jsonify({
            'success': True,
            'cover_letter': new_cover_letter
        })

    except Exception as e:
        return jsonify({
            'error': str(e)
        }), 500


@app.route('/download/<format>', methods=['POST'])
def download_cover_letter(format):
    """Download cover letter in specified format"""
    try:
        data = request.get_json()
        text = data.get('text', '')

        if not text:
            return jsonify({'error': 'No text provided'}), 400

        # Generate filename with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        if format == 'txt':
            filename = f'cover_letter_{timestamp}.txt'
            return send_file(
                io.BytesIO(text.encode('utf-8')),
                mimetype='text/plain',
                as_attachment=True,
                download_name=filename
            )

        elif format == 'docx':
            filename = f'cover_letter_{timestamp}.docx'
            docx_file = generate_docx(text)
            return send_file(
                docx_file,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=filename
            )

        elif format == 'pdf':
            filename = f'cover_letter_{timestamp}.pdf'
            pdf_file = generate_pdf(text)
            return send_file(
                pdf_file,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=filename
            )

        else:
            return jsonify({'error': 'Unsupported format'}), 400

    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    print("=" * 60)
    print("🎯 Cover Letter Generator Server")
    print("=" * 60)
    print(f"AI Provider: {AI_PROVIDER}")
    print(f"AI Configured: {bool(ANTHROPIC_API_KEY or OPENAI_API_KEY)}")
    if not (ANTHROPIC_API_KEY or OPENAI_API_KEY):
        print("\n⚠️  WARNING: No AI API key configured!")
        print("Set ANTHROPIC_API_KEY or OPENAI_API_KEY environment variable")
        print("The tool will use basic templates without AI")
    print("\n🚀 Server starting on http://localhost:8080")
    print("=" * 60)

    app.run(host='0.0.0.0', port=8080, debug=True)
